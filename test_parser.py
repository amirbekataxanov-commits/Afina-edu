from docx import Document
import re
import uuid

class Option:
    def __init__(self, letter, option, is_correct: bool):
        self.letter = letter
        self.option = option
        self.is_correct = is_correct

    def __repr__(self):
        return f"{self.letter}) {self.option}{' *' if self.is_correct else ''}"

class Question:
    def __init__(self, question: str, options: list[Option]):
        self.question = question
        self.options = options

    def __repr__(self):
        return f"Q: {self.question}\n" + "\n".join(str(opt) for opt in self.options)
    



def parse_docx_tests(file_path) -> list[Question]:
    document = Document(file_path)
    tests = []

    current_question_text = None
    current_options = []

    option_pattern = re.compile(r"^([A-D])\)\s*(.+?)(\s*\*)?$")
    question_pattern = re.compile(r"^\d+\.\s*(.+)")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        question_match = question_pattern.match(text)
        if question_match:
            if current_question_text:
                # Save the previous question
                tests.append(Question(current_question_text, current_options))
            current_question_text = question_match.group(1)
            current_options = []

        else:
            option_match = option_pattern.match(text)
            if option_match:
                letter = option_match.group(1)
                option_text = option_match.group(2).strip()
                is_correct = option_match.group(3) is not None
                current_options.append(Option(letter, option_text, is_correct))

    # Append the last question
    if current_question_text:
        tests.append(Question(current_question_text, current_options))

    return tests

def parse_and_save(test_file_path) -> str: 
    tests = parse_docx_tests(test_file_path)
    dest = f"temp_tests/{str(uuid.uuid4())}.txt"
    with open(dest, "w", encoding="utf-8") as f:
         for test in tests:
            f.write(f"SAVOL: {test.question}\n")
            for opt in test.options:
                f.write(f"{opt.letter}) {opt.option}{' *' if opt.is_correct else ''}\n")
            f.write("\n")
    
    return dest


def parse_single_test(test_text: str) -> Question | None:
    """
    Parse a single test from a string and return a Question model.
    
    Example input:
    "1. Movarounnahrda Temuriylar davlati o‘rnida qaysi sulola hokimiyatni egalladi?\n
    A) Safaviylar\n
    B) Ashtarxoniylar\n
    C) Shayboniylar *\n
    D) Mang‘itlar"
    """
    # Patterns
    question_pattern = re.compile(r"^\d+\.\s*(.+)")
    option_pattern = re.compile(r"^([A-D])\)\s*(.+?)(\s*\*)?$")

    # Split lines and clean
    lines = [line.strip() for line in test_text.strip().splitlines() if line.strip()]
    if not lines:
        return None

    # Extract question
    question_match = question_pattern.match(lines[0])
    if not question_match:
        return None  # invalid format

    question_text = question_match.group(1)
    options = []

    # Extract options
    for line in lines[1:]:
        option_match = option_pattern.match(line)
        if option_match:
            letter = option_match.group(1)
            option_text = option_match.group(2).strip()
            is_correct = option_match.group(3) is not None
            options.append(Option(letter, option_text, is_correct))

    if not options:
        return None  # no options found

    return Question(question_text, options)



class FillBlanksTest:
    def __init__(self, text, answers):
        self.text = text
        self.answers = answers  # list[FillBlanksTestAnswer]

class FillBlanksTestAnswer:
    def __init__(self, number, answer):
        self.number = number
        self.answer = answer

def parse_docx_fill_blanks(file_path) -> list[FillBlanksTest]:
    document = Document(file_path)
    tests = []

    current_text = None
    current_answers = []

    question_pattern = re.compile(r"^\d+\)")        # 1) ...
    answer_pattern = re.compile(r"^(\d+)\.\s*(.+)") # 1. answer

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if question_pattern.match(text):
            if current_text:
                tests.append(FillBlanksTest(current_text, current_answers))
            current_text = text[text.find(')') + 1:].strip()
            current_answers = []

        else:
            match = answer_pattern.match(text)
            if match:
                number = match.group(1).strip()
                answer = match.group(2).strip()
                current_answers.append(FillBlanksTestAnswer(number, answer))

    # append the last parsed test
    if current_text:
        tests.append(FillBlanksTest(current_text, current_answers))

    return tests


# TEST SECTION
# if __name__ == "__main__":
#     file_path = r"D:\PROJECTES\AfinaBot\src\fakedata\1-mavzu.docx"
#     tests = parse_docx_fill_blanks(file_path=file_path)
#     print(tests)